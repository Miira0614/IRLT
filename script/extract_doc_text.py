# -*- coding: utf-8 -*-
"""
规范书文本提取工具
从 .doc 文件中提取纯文本内容，支持多种方式：
1. win32com (Word/WPS COM automation) - 推荐，Windows下最可靠
2. python-docx (需要先手动另存为 .docx)
3. olefile 直接解析（无需Office，但格式保真度较低）
"""
import os
import sys
import re


def extract_via_com(doc_path: str) -> str:
    """
    通过 Windows COM 自动化提取 .doc 文本
    支持 Microsoft Word 和 WPS Office
    """
    try:
        import win32com.client
    except ImportError:
        raise ImportError("请先安装 pywin32: pip install pywin32")

    word = None
    doc = None
    try:
        # 先尝试 Microsoft Word
        try:
            word = win32com.client.Dispatch("Word.Application")
        except Exception:
            # 尝试 WPS Office
            try:
                word = win32com.client.Dispatch("WPS.Application")
            except Exception:
                try:
                    word = win32com.client.Dispatch("KWPS.Application")
                except Exception:
                    raise RuntimeError(
                        "无法启动 Word 或 WPS。请确保已安装 Microsoft Word 或 WPS Office。"
                    )

        word.Visible = False
        word.DisplayAlerts = 0

        abs_path = os.path.abspath(doc_path)
        doc = word.Documents.Open(abs_path, ReadOnly=True)

        # 提取全部文本
        text = doc.Content.Text

        doc.Close()
        word.Quit()

        return text

    except Exception as e:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        raise RuntimeError(f"COM提取文本失败: {str(e)}")


def extract_via_olefile(doc_path: str) -> str:
    """
    通过 olefile 直接从 OLE2 复合文档中提取文本
    不需要安装 Office，但可能丢失格式和表格结构
    """
    try:
        import olefile
    except ImportError:
        raise ImportError("请先安装 olefile: pip install olefile")

    ole = olefile.OleFileIO(doc_path)

    # 方法1: 尝试读取 WordDocument 流
    text_parts = []

    if ole.exists("WordDocument"):
        data = ole.openstream("WordDocument").read()
        # 尝试提取 Unicode 文本
        try:
            # Word 文档的文本通常以 UTF-16LE 编码存储在特定位置
            # 这里做简单处理：提取所有可读的 Unicode 字符
            decoded = data.decode("utf-16-le", errors="ignore")
            # 过滤控制字符
            text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', decoded)
            text_parts.append(text)
        except Exception:
            pass

    # 方法2: 读取 1Table 或 0Table 流
    for stream_name in ["1Table", "0Table"]:
        if ole.exists(stream_name):
            try:
                data = ole.openstream(stream_name).read()
                decoded = data.decode("utf-16-le", errors="ignore")
                text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', decoded)
                text_parts.append(text)
            except Exception:
                pass

    ole.close()

    if not text_parts:
        raise RuntimeError("无法从 .doc 文件中提取文本。建议使用 COM 方式或先转换为 .docx。")

    return "\n\n".join(text_parts)


def extract_via_libreoffice(doc_path: str, output_dir: str = None) -> str:
    """
    通过 LibreOffice 命令行转换为文本
    需要安装 LibreOffice
    """
    import subprocess
    import tempfile

    if output_dir is None:
        output_dir = tempfile.mkdtemp()

    abs_path = os.path.abspath(doc_path)

    # 尝试找到 LibreOffice
    libreoffice_paths = [
        "soffice",
        "libreoffice",
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]

    cmd = None
    for path in libreoffice_paths:
        try:
            subprocess.run([path, "--version"], capture_output=True, timeout=5)
            cmd = path
            break
        except Exception:
            continue

    if cmd is None:
        raise RuntimeError("未找到 LibreOffice。请安装 LibreOffice 或使用 COM 方式。")

    # 转换为 txt
    subprocess.run(
        [cmd, "--headless", "--convert-to", "txt:Text",
         "--outdir", output_dir, abs_path],
        timeout=60, check=True
    )

    # 找到生成的 txt 文件
    base_name = os.path.splitext(os.path.basename(doc_path))[0]
    txt_path = os.path.join(output_dir, base_name + ".txt")

    if not os.path.exists(txt_path):
        raise RuntimeError(f"LibreOffice 转换失败，未生成文本文件: {txt_path}")

    with open(txt_path, 'r', encoding='utf-8') as f:
        text = f.read()

    return text


def extract_via_docx_conversion(doc_path: str) -> str:
    """
    通过 COM 将 .doc 转为 .docx，然后用 python-docx 读取
    """
    try:
        import win32com.client
    except ImportError:
        raise ImportError("请先安装 pywin32: pip install pywin32")

    try:
        from docx import Document
    except ImportError:
        raise ImportError("请先安装 python-docx: pip install python-docx")

    import tempfile

    word = None
    doc = None
    try:
        try:
            word = win32com.client.Dispatch("Word.Application")
        except Exception:
            try:
                word = win32com.client.Dispatch("WPS.Application")
            except Exception:
                raise RuntimeError("无法启动 Word 或 WPS。")

        word.Visible = False
        word.DisplayAlerts = 0

        abs_path = os.path.abspath(doc_path)
        doc = word.Documents.Open(abs_path, ReadOnly=True)

        # 保存为 .docx
        temp_dir = tempfile.mkdtemp()
        docx_path = os.path.join(temp_dir, "temp_spec.docx")
        doc.SaveAs2(docx_path, FileFormat=16)  # 16 = wdFormatDocumentDefault (docx)

        doc.Close()
        word.Quit()

        # 用 python-docx 读取
        docx_doc = Document(docx_path)
        paragraphs = [p.text for p in docx_doc.paragraphs]

        # 也读取表格内容
        for table in docx_doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                paragraphs.append(row_text)

        # 清理临时文件
        try:
            os.remove(docx_path)
            os.rmdir(temp_dir)
        except Exception:
            pass

        return "\n".join(paragraphs)

    except Exception as e:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        raise RuntimeError(f"COM转换提取失败: {str(e)}")


def extract_text(doc_path: str, method: str = "auto") -> str:
    """
    提取 .doc 文件的文本内容

    Args:
        doc_path: .doc 文件路径
        method: 提取方式
            - "auto": 自动选择最佳方式（按优先级尝试）
            - "com": 强制使用 COM (Word/WPS)
            - "docx_convert": COM转为docx后用python-docx读取（保留表格结构）
            - "libreoffice": 使用 LibreOffice 转换
            - "olefile": 直接解析 OLE2 格式

    Returns:
        提取的文本内容
    """
    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"文件不存在: {doc_path}")

    methods = {
        "docx_convert": extract_via_docx_conversion,
        "com": extract_via_com,
        "libreoffice": extract_via_libreoffice,
        "olefile": extract_via_olefile,
    }

    if method != "auto":
        if method not in methods:
            raise ValueError(f"不支持的提取方式: {method}。可选: {list(methods.keys())}")
        return methods[method](doc_path)

    # 自动模式：按优先级尝试
    errors = []
    for method_name in ["docx_convert", "com", "libreoffice", "olefile"]:
        try:
            print(f"[INFO] 尝试使用 {method_name} 方式提取...")
            text = methods[method_name](doc_path)
            if text and len(text.strip()) > 100:
                print(f"[OK] {method_name} 方式提取成功，共 {len(text)} 字符")
                return text
            else:
                errors.append(f"{method_name}: 提取内容过短（{len(text) if text else 0} 字符）")
        except Exception as e:
            errors.append(f"{method_name}: {str(e)[:100]}")
            continue

    raise RuntimeError(
        f"所有提取方式均失败:\n" + "\n".join(f"  - {e}" for e in errors)
    )


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="从 .doc 规范书中提取文本")
    parser.add_argument("doc_path", help=".doc 文件路径")
    parser.add_argument("-o", "--output", default=None, help="输出文本文件路径")
    parser.add_argument("-m", "--method", default="auto",
                        choices=["auto", "com", "docx_convert", "libreoffice", "olefile"],
                        help="提取方式")
    args = parser.parse_args()

    try:
        text = extract_text(args.doc_path, method=args.method)
        print(f"\n[OK] 成功提取 {len(text)} 字符")

        if args.output:
            os.makedirs(os.path.dirname(args.output) or ".", exist_ok=True)
            with open(args.output, 'w', encoding='utf-8') as f:
                f.write(text)
            print(f"[OK] 已保存到: {args.output}")
        else:
            # 预览前2000字符
            print("\n--- 文本预览（前2000字符）---")
            print(text[:2000])
            if len(text) > 2000:
                print(f"\n... (共 {len(text)} 字符，已截断)")

    except Exception as e:
        print(f"[ERROR] {str(e)}")
        sys.exit(1)
